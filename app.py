import streamlit as st
from google import genai
import fitz  # PyMuPDF
from docx import Document
import io
import time

# --- UI CONFIGURATION ---
st.set_page_config(page_title="Physics Board Bot 2026", page_icon="🍎")
st.title("🍎 Physics Board Exam Expert")
st.markdown("---")

# --- SIDEBAR: API CONFIG ---
with st.sidebar:
    st.header("Settings")
    user_api_key = st.text_input("Gemini API Key daalein:", type="password")
    st.info("Aapki key surakshit hai aur sirf is session ke liye use hogi.")

# --- MAIN APP LOGIC ---
uploaded_file = st.file_uploader("Chapter PDF Upload Karein", type="pdf")

if uploaded_file and user_api_key:
    # 1. Extract Text from PDF
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text()

    st.success(f"PDF Uploaded! Total Pages: {len(doc)}")

    if st.button("Deep Notes & Question Bank Taiyaar Karein"):
        try:
            client = genai.Client(api_key=user_api_key)
            
            # --- PROGRESS BAR ---
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            # Master Prompt Logic
            master_prompt = f"""
            You are a Senior Physics Board Examiner. Create "Deep Text-Only Notes" for this text.
            
            STRICT RULES:
            - NO IMAGES & NO LaTeX: Use plain text only (e.g., V = I x R, H = I squared R t).
            - NO STARS: Use '-' for bullet points.
            - 2026 PATTERN: Add 1 Case-Based, 2 Assertion-Reason, and 3 Competency questions after each topic.
            - SYLLABUS: Cover all Class 10 Board topics (Current, Potential, Ohm's Law, Resistance, Heating Effect, Power).
            - DEPTH: Explain the conceptual 'Why' for everything.

            Text: {full_text[:15000]} 
            """
            
            status_text.text("AI Thinking... Board Pattern analyze ho raha hai.")
            response = client.models.generate_content(model="gemini-2.0-flash", contents=master_prompt)
            progress_bar.progress(100)
            
            # --- DISPLAY RESULTS ---
            st.markdown("### 📝 Generated Deep Notes")
            st.write(response.text)

            # --- EXPORT TO WORD ---
            word_doc = Document()
            word_doc.add_heading('Physics Deep Notes & 2026 Question Bank', 0)
            word_doc.add_paragraph(response.text)
            
            # Save to memory
            bio = io.BytesIO()
            word_doc.save(bio)
            
            st.download_button(
                label="📥 Final Word File Download Karein",
                data=bio.getvalue(),
                file_name="Physics_Board_Expert_Notes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except Exception as e:
            st.error(f"Error: {e}")

else:
    st.info("Aage badhne ke liye apni API Key daalein aur PDF upload karein.")